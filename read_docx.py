import docx
import io

try:
    doc = docx.Document(r'c:\Users\datmk\OneDrive\Desktop\Workspace\NemoClaw\apps\client-react\src\components\Tổ hợp xét tuyển đại học năm 2026.docx')
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    tables_data = []
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            tables_data.append(' | '.join(row_data))

    with io.open('extracted_data.txt', 'w', encoding='utf-8') as f:
        f.write("--- PARAGRAPHS ---\n")
        f.write('\n'.join(paragraphs))
        f.write("\n--- TABLES ---\n")
        f.write('\n'.join(tables_data))
except Exception as e:
    with io.open('extracted_data.txt', 'w', encoding='utf-8') as f:
        f.write(f"Error: {e}")
